"""Resume file text extraction."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


class UnsupportedFileTypeError(ValueError):
    """Raised when a user uploads an unsupported document type."""


class EmptyResumeError(ValueError):
    """Raised when no readable text can be extracted from a resume."""


_PDF_MAGIC = b"%PDF"
_DOCX_MAGIC = b"PK\x03\x04"


class ResumeParser:
    """Extract raw text from supported resume formats."""

    supported_extensions = {".pdf", ".docx"}

    def extract_text(self, file_path: Path) -> str:
        """Extract normalized text from a PDF or DOCX resume."""
        extension = file_path.suffix.lower()
        if extension == ".pdf":
            text = self._extract_pdf(file_path)
        elif extension == ".docx":
            text = self._extract_docx(file_path)
        else:
            raise UnsupportedFileTypeError("Only PDF and DOCX resumes are supported.")

        normalized = _normalize_text(text)
        if not normalized:
            raise EmptyResumeError("Could not extract readable text from the resume.")
        return normalized

    @classmethod
    def validate_magic_bytes(cls, content: bytes, extension: str) -> None:
        """Raise UnsupportedFileTypeError if the file header doesn't match the extension."""
        if extension == ".pdf" and not content.startswith(_PDF_MAGIC):
            raise UnsupportedFileTypeError(
                "Uploaded file does not appear to be a valid PDF."
            )
        if extension == ".docx" and not content.startswith(_DOCX_MAGIC):
            raise UnsupportedFileTypeError(
                "Uploaded file does not appear to be a valid DOCX."
            )

    @staticmethod
    def _extract_pdf(file_path: Path) -> str:
        import fitz

        parts: list[str] = []
        with fitz.open(file_path) as document:
            for page in document:
                parts.append(page.get_text("text"))
        return "\n".join(parts)

    @staticmethod
    def _extract_docx(file_path: Path) -> str:
        document = Document(file_path)
        parts: list[str] = []

        for section in document.sections:
            for para in section.header.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for para in section.footer.paragraphs:
                if para.text.strip():
                    parts.append(para.text)

        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        for txbx in document.element.body.iter(qn("w:txbxContent")):
            for para in txbx.iter(qn("w:p")):
                text = "".join(run.text or "" for run in para.iter(qn("w:t")))
                if text.strip():
                    parts.append(text.strip())

        return "\n".join(parts)


def _normalize_text(text: str) -> str:
    """Collapse intra-line whitespace and strip blank lines."""
    lines = (" ".join(line.split()) for line in text.splitlines())
    return "\n".join(line for line in lines if line).strip()
